import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether, HRFlowable
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

doc_title = "Google Gemini Research: AI Visibility & Generative Engine Optimization (GEO)"
source_url = "https://share.gemini.google/WIKDgMGnPTVL"
docx_path = r"c:\AIQualityHQ\Google_Gemini_AI_Visibility_Research.docx"
pdf_path = r"c:\AIQualityHQ\Google_Gemini_AI_Visibility_Research.pdf"

# --- 1. GENERATE WORD (.DOCX) FILE ---
def create_word_docx():
    doc = Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # Styles Setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Helper functions for Docx styling
    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D) # Navy Blue
        p.paragraph_format.space_after = Pt(4)
        return p

    def add_subtitle(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
        p.paragraph_format.space_after = Pt(18)
        return p

    def add_heading1(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        return p

    def add_heading2(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        return p

    def add_paragraph(text, bold_prefix=""):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.bold = True
        p.add_run(text)
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.bold = True
        p.add_run(text)
        return p

    def set_cell_background(cell, fill_hex):
        tcPr = cell._element.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    # Document Header
    add_title("Google Gemini Research: AI Visibility & GEO")
    add_subtitle(f"Source: {source_url} | Extracted: July 30, 2026")

    # Section 1
    add_heading1("1. Core Research Findings & Empirical Frameworks")
    add_bullet(" Generative search marks a fundamental transition from traditional keyword-based indexing to semantic retrieval via Retrieval-Augmented Generation (RAG).", "The RAG Shift:")
    add_bullet(" On-page prompts and instructions cannot dictate what a model says about a brand. The RAG pipeline independently queries the web for real-time sources. AI discoverability requires technical crawlability, structured machine-readable entities, and external third-party citations.", "The Prompt Engineering Fallacy:")
    
    add_heading2("Empirical GEO Optimization Framework (Princeton Study)")
    add_paragraph("Position-Adjusted Word Count (PAWC) serves as the standard empirical metric for measuring generative engine citation visibility.")
    add_bullet(" Adding direct, authoritative quotes increases PAWC visibility by +41%.", "Incorporate Expert Quotations:")
    add_bullet(" Replacing vague qualitative text with verified statistics increases PAWC by +31%.", "Include Quantitative Data:")
    add_bullet(" Ensuring high readability and clear syntax yields a +28% improvement.", "Optimize Content Fluency:")
    add_bullet(" Referencing reputable third-party studies leads to a +28% increase.", "Add Inline Citations:")
    add_bullet(" Repetitive keyword placement results in an -8.0% citation drop as models treat it as low-quality spam.", "Avoid Keyword Stuffing:")
    add_bullet(" RAG models often prioritize niche, structurally extractable passages or authoritative discussions over top-ranking organic Google pages. Rank 5 organic content frequently outperforms Rank 1 in LLM recommendation rate.", "SERP Equalization:")

    # Section 2
    add_heading1("2. Target Reader Personas")
    
    add_heading2("Persona A: Technical SEO Director (Enterprise B2B SaaS)")
    add_bullet("Protect organic search market share and establish LLM crawlability standards.", "Primary Objectives: ")
    add_bullet("Skeptical of superficial marketing advice. Needs precise technical specifications for schema, browser audits, crawler behavior, and JavaScript hydration.", "Pain Points: ")
    add_bullet("Code-level documentation, raw vs. rendered DOM comparisons, JSON-LD, and crawler user-agent configuration.", "Preferred Depth: ")

    add_heading2("Persona B: Head of Growth / Marketing (Scaling Startup)")
    add_bullet("Drive qualified referral traffic and ensure brand inclusion in ChatGPT and Perplexity buyer recommendations.", "Primary Objectives: ")
    add_bullet("Frustrated that competitors win ChatGPT brand recommendations; needs clear attribution frameworks and measurable ROI.", "Pain Points: ")
    add_bullet("Executive summaries, competitive benchmarking tables, ROI attribution strategies, and strategic workflows.", "Preferred Depth: ")

    add_heading2("Persona C: Founder / Indie Hacker (Early-Stage Tech / Software)")
    add_bullet("Resolve search engine indexation issues quickly, leverage AI coding agents for development speed, and drive early adoption.", "Primary Objectives: ")
    add_bullet("Limited engineering bandwidth and budget; wants to avoid complex sales calls or opaque pricing.", "Pain Points: ")
    add_bullet("Step-by-step developer guides, ready-to-paste AI coding workspace prompts, and direct tool pricing comparisons.", "Preferred Depth: ")

    add_heading2("Persona D: Digital Agency Strategist (Search & Marketing Agencies)")
    add_bullet("Upsell GEO and AI visibility tracking retainers to existing clients.", "Primary Objectives: ")
    add_bullet("Needs client-ready white-label reporting dashboard systems and multi-brand tracking workflows.", "Pain Points: ")
    add_bullet("Pitch decks, reporting standards, and feature matrices comparing visibility dashboard solutions.", "Preferred Depth: ")

    # Section 3
    add_heading1("3. Keyword Research & Semantic Entity Matrix")
    add_bullet("AI visibility (Commercial / Informational) | Entities: Generative Engine Optimization, LLM SEO, RAG Retrieval | Surfaces: Google Search, Perplexity", "Primary Keyword: ")
    add_bullet("Generative Engine Optimization, LLM SEO, AI search optimization, AI citation tracking, brand visibility tracking platform | Entities: Knowledge Graph, Schema.org, Vector Embeddings, TrueDR", "Secondary Keywords: ")
    add_bullet("tracking brand mentions in ChatGPT and Perplexity, why ChatGPT recommends competitors, fix website JavaScript blindness for AI crawlers | Entities: OAI-SearchBot, PerplexityBot, Rendered DOM Diff", "Long-Tail Keywords: ")
    add_bullet("How do AI search engines pick sources?, Why isn't my site cited by ChatGPT?, What is the difference between SEO and GEO? | Entities: Passage Retrieval, Query Fan-out, PAWC", "Question Keywords: ")

    add_heading2("Conversational Buying-Intent Prompts")
    add_bullet("\"What are the top AI visibility tracking platforms for B2B SaaS companies in 2026?\"")
    add_bullet("\"Compare Peec AI, Profound, and VerifiedDR for tracking brand citations across ChatGPT and Perplexity.\"")
    add_bullet("\"Which software platforms identify why ChatGPT recommends my competitors instead of my product?\"")
    add_bullet("\"How can I audit my web page to ensure OpenAI's OAI-SearchBot and PerplexityBot can read my HTML?\"")

    # Section 4
    add_heading1("4. Competitive Landscape & Tool Comparison Matrix")
    add_paragraph("The monitoring-to-optimization spectrum spans four distinct capability tiers: Level 1 (Basic Mention Tracking), Level 2 (Analytics & Share of Voice), Level 3 (Optimization Recommendations), and Level 4 (Closed-Loop Agentic Execution).")

    # Table in Word
    table_data = [
        ["Platform", "Target Audience", "Pricing", "Tracking Method", "Key Strengths", "Best Use Case"],
        ["VerifiedDR", "Founders, SEOs, agencies, dev teams", "Flexible tiers (Pro/Agency)", "Local browser execution (47 checks)", "Audits 5+ bots, raw/rendered HTML diff, TrueDR, Claude/Cursor prompts", "Teams bridging technical audits with direct agentic code fixes"],
        ["Profound", "Enterprise brands", "Custom Enterprise", "API-based multi-model prompt runs", "SOC 2 compliance, governance, executive dashboards, historical archives", "Fortune 500 brands needing high-level share-of-voice reporting"],
        ["Peec AI", "Startups, mid-market agencies", "$95/mo - $245/mo", "Daily automated prompt runs", "Clean UI, unlimited seats, AI Shopping Analytics for e-commerce", "Marketing teams tracking shopping placements & daily prompt ranks"],
        ["ZipTie.dev", "SEOs, content strategists", "~$79/mo", "Real UI web crawling & query discovery", "Simulates user UI behaviors, generates auto queries from URLs", "SEO teams looking to locate content gaps based on real UI sessions"],
        ["Otterly.AI", "Small businesses, solo marketers", "~$29/mo", "Basic multi-model monitoring", "Low cost, quick setup, simple dashboards", "Budget-conscious teams needing basic brand alerts in ChatGPT"],
        ["Scrunch AI", "Data analysts, enterprise agencies", "$83/mo - $417+/mo", "Cohort measurement & analytics engine", "Cohort breakdown, metric customization", "Data-driven agencies running complex multi-brand comparisons"],
        ["Hall", "Brand managers, PR specialists", "$199/mo - $499+/mo", "Always-on continuous query monitoring", "Monitors citation sentiment, citation drift, crawler activity", "Brand protection teams tracking ongoing sentiment drift"]
    ]

    t = doc.add_table(rows=len(table_data), cols=6)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False

    col_widths = [Inches(1.0), Inches(1.1), Inches(0.9), Inches(1.1), Inches(1.4), Inches(1.3)]
    for i, row in enumerate(t.rows):
        for j, cell in enumerate(row.cells):
            cell.width = col_widths[j]
            cell.paragraphs[0].text = table_data[i][j]
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            cell.paragraphs[0].paragraph_format.line_spacing = 1.05
            p_run = cell.paragraphs[0].runs[0]
            p_run.font.size = Pt(8.5)
            if i == 0:
                p_run.font.bold = True
                p_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                set_cell_background(cell, "1A365D")
            else:
                if i % 2 == 1:
                    set_cell_background(cell, "F7FAFC")
                else:
                    set_cell_background(cell, "EDF2F7")

    # Section 5
    add_heading1("5. Technical Analysis of VerifiedDR")
    add_paragraph("Combines on-page technical diagnostics (via a local 47-check browser audit) with multi-engine query tracking to identify indexing blockages and resolve JavaScript blindness. Outputs developer-ready prompts for Claude Code or Cursor.")
    
    add_heading2("The 47-Check Browser Audit Engine")
    add_bullet(" Inspects robots.txt configurations to verify permissions for GPTBot, OAI-SearchBot, ClaudeBot, PerplexityBot, and Google-Extended.", "1. AI Crawler Access Verification:")
    add_bullet(" Performs raw HTML vs. rendered DOM difference analysis to ensure client-side hydration doesn't render content invisible to scraping bots.", "2. JavaScript Blindness Detection:")
    add_bullet(" Validates JSON-LD schema configurations for Organization, Product, FAQPage, HowTo, Article, and date tags.", "3. Structured Data Validation:")
    add_bullet(" Checks for valid llms.txt and llms-full.txt markdown files in the domain root.", "4. llms.txt Machine File Verification:")
    add_bullet(" Calculates refined domain authority score by cross-referencing DR with verified organic traffic, indexation rate, and search trust signals.", "5. TrueDR Authority Verification:")

    # Section 6
    add_heading1("6. Optimization Checklists")
    
    add_heading2("AI Citation & GEO Optimization Checklist")
    add_bullet("Incorporate direct, attribute-based quotes from recognized subject matter experts (+41% PAWC impact).")
    add_bullet("Replace qualitative assertions with statistics, numbers, and concrete data points (+31% PAWC impact).")
    add_bullet("Focus on direct syntax, clear structures, and direct-answer openings (+28% PAWC impact).")
    add_bullet("Reference external studies, whitepapers, and documentation (+28% PAWC impact).")
    add_bullet("Remove repetitive keyword lists and forced matching phrases (-8% PAWC penalty).")

    add_heading2("Technical SEO & LLM Crawlability Checklist")
    add_bullet("Ensure robots.txt explicitly allows crawler bots (OAI-SearchBot, PerplexityBot, ClaudeBot, Google-Extended).")
    add_bullet("Perform DOM-to-raw-HTML diff checks to verify text loads without client-side scripts.")
    add_bullet("Create and host a clear markdown site summary at yourdomain.com/llms.txt.")
    add_bullet("Format content using clean semantic HTML5 (<article>, <section>, <h1>–<h3>) for RAG passage chunking.")

    # Section 7 & 8 & 9 & 10
    add_heading1("7. Strategic Editorial Blueprint")
    add_paragraph("Title: Why Prompt Engineering Won't Save Your AI Visibility (And What Actually Works in 2026)")
    add_paragraph("Target Length: 3,500 – 4,500 words.")

    add_heading1("8. Fact Verification Matrix")
    add_bullet(" Adding quotations & stats increases visibility +31% to +41% (Verified Fact - Princeton Study).")
    add_bullet(" Keyword stuffing reduces AI answer visibility by -8.0% (Verified Fact).")
    add_bullet(" Google AI Overviews overlap 76.1% with top 10 organic SERP ranks (Verified Fact).")
    add_bullet(" Disallowing OAI-SearchBot in robots.txt blocks ChatGPT Search citations (Verified Fact).")

    add_heading1("9. Strategic Implementation Roadmap")
    add_bullet("Audit robots.txt, analyze JS blindness, deploy root llms.txt, validate JSON-LD.", "Phase 1: Technical Foundation Audit (Days 1–14): ")
    add_bullet("Add expert quotes and stats, restructure with direct-answer openings (AEO blocks).", "Phase 2: Content Restructuring & GEO (Days 15–45): ")
    add_bullet("Track neutral prompts, export diagnostic results into developer prompts for Claude Code/Cursor.", "Phase 3: Ongoing Buying-Intent Monitoring: ")

    doc.save(docx_path)
    print(f"Docx saved to {docx_path}")

# --- 2. GENERATE PDF FILE ---
def create_pdf():
    doc = SimpleDocTemplate(
        pdf_path,
        pagesize=letter,
        leftMargin=0.75*inch,
        rightMargin=0.75*inch,
        topMargin=0.75*inch,
        bottomMargin=0.75*inch
    )

    styles = getSampleStyleSheet()
    
    # Custom Palette
    c_primary = colors.HexColor("#1A365D")
    c_secondary = colors.HexColor("#2B6CB0")
    c_text = colors.HexColor("#2D3748")

    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=20,
        leading=24,
        textColor=c_primary,
        spaceAfter=4
    )
    
    sub_style = ParagraphStyle(
        'DocSubTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Oblique',
        fontSize=10,
        leading=13,
        textColor=colors.HexColor("#718096"),
        spaceAfter=14
    )

    h1_style = ParagraphStyle(
        'SectionH1',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=14,
        leading=18,
        textColor=c_primary,
        spaceBefore=14,
        spaceAfter=6,
        keepWithNext=True
    )

    h2_style = ParagraphStyle(
        'SectionH2',
        parent=styles['Heading3'],
        fontName='Helvetica-Bold',
        fontSize=11,
        leading=15,
        textColor=c_secondary,
        spaceBefore=10,
        spaceAfter=4,
        keepWithNext=True
    )

    body_style = ParagraphStyle(
        'BodyTextCustom',
        parent=styles['BodyText'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13.5,
        textColor=c_text,
        spaceAfter=6
    )

    bullet_style = ParagraphStyle(
        'BulletCustom',
        parent=styles['BodyText'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13.5,
        textColor=c_text,
        leftIndent=15,
        firstLineIndent=-10,
        spaceAfter=4
    )

    story = []

    story.append(Paragraph("Google Gemini Research: AI Visibility & GEO Framework", title_style))
    story.append(Paragraph(f"Source: {source_url} | Extracted & Prepared: July 30, 2026", sub_style))
    story.append(HRFlowable(width="100%", thickness=1, color=c_secondary, spaceAfter=12))

    # Section 1
    story.append(Paragraph("1. Core Research Findings & Empirical Frameworks", h1_style))
    story.append(Paragraph("<b>• The RAG Shift:</b> Generative search marks a fundamental transition from keyword-based indexing to semantic retrieval via Retrieval-Augmented Generation (RAG).", bullet_style))
    story.append(Paragraph("<b>• The Prompt Engineering Fallacy:</b> On-page prompts cannot dictate LLM outputs. AI discoverability requires technical crawlability, structured machine-readable entities, and external third-party citations.", bullet_style))

    story.append(Paragraph("Empirical GEO Optimization Framework (Princeton Study)", h2_style))
    story.append(Paragraph("Position-Adjusted Word Count (PAWC) is the standard metric for measuring generative engine citation visibility.", body_style))
    story.append(Paragraph("<b>• Expert Quotations:</b> Direct, authoritative quotes increase PAWC visibility by <b>+41%</b>.", bullet_style))
    story.append(Paragraph("<b>• Quantitative Data:</b> Replacing vague text with statistics increases PAWC by <b>+31%</b>.", bullet_style))
    story.append(Paragraph("<b>• Content Fluency:</b> High readability yields a <b>+28%</b> improvement.", bullet_style))
    story.append(Paragraph("<b>• Inline Citations:</b> Referencing reputable third-party studies adds <b>+28%</b> visibility.", bullet_style))
    story.append(Paragraph("<b>• Avoid Keyword Stuffing:</b> Causes an <b>-8.0%</b> citation drop due to spam classification.", bullet_style))
    story.append(Paragraph("<b>• SERP Equalization:</b> Rank 5 organic content frequently outperforms Rank 1 in LLM recommendation rate.", bullet_style))

    # Section 2
    story.append(Paragraph("2. Target Reader Personas", h1_style))
    story.append(Paragraph("<b>• Persona A: Technical SEO Director:</b> Focuses on enterprise B2B SaaS crawlability, JSON-LD, JS hydration, and raw/rendered DOM diffs.", bullet_style))
    story.append(Paragraph("<b>• Persona B: Head of Growth:</b> Focuses on ChatGPT/Perplexity buyer recommendations, ROI attribution, and competitive share of voice.", bullet_style))
    story.append(Paragraph("<b>• Persona C: Founder / Indie Hacker:</b> Needs step-by-step developer guides and AI coding prompts (Claude Code/Cursor).", bullet_style))
    story.append(Paragraph("<b>• Persona D: Agency Strategist:</b> Requires client-ready white-label reporting dashboards and multi-brand workflows.", bullet_style))

    # Section 3
    story.append(Paragraph("3. Keyword Research & Semantic Entity Matrix", h1_style))
    story.append(Paragraph("<b>• Primary Keyword:</b> <i>AI visibility</i> (Entities: Generative Engine Optimization, LLM SEO, RAG Retrieval)", bullet_style))
    story.append(Paragraph("<b>• Key Technical Entities:</b> Knowledge Graph, Schema.org, Vector Embeddings, TrueDR, Passage Retrieval", bullet_style))
    story.append(Paragraph("<b>• Buying-Intent Prompts:</b> \"Top AI visibility tracking platforms for B2B SaaS in 2026\", \"Compare Peec AI, Profound, and VerifiedDR\".", bullet_style))

    # Section 4 Table
    story.append(Paragraph("4. Competitive Landscape & Tool Comparison", h1_style))
    
    table_raw_data = [
        [Paragraph("<b>Platform</b>", body_style), Paragraph("<b>Target</b>", body_style), Paragraph("<b>Pricing</b>", body_style), Paragraph("<b>Methodology</b>", body_style), Paragraph("<b>Key Strengths</b>", body_style)],
        [Paragraph("<b>VerifiedDR</b>", body_style), Paragraph("Founders, SEOs, Devs", body_style), Paragraph("Pro / Agency", body_style), Paragraph("Local 47-check audit", body_style), Paragraph("Audits 5+ bots, raw/rendered HTML diff, TrueDR, Cursor/Claude prompts", body_style)],
        [Paragraph("<b>Profound</b>", body_style), Paragraph("Enterprise", body_style), Paragraph("Custom", body_style), Paragraph("API prompt runs", body_style), Paragraph("SOC 2, governance, executive dashboards, historical archives", body_style)],
        [Paragraph("<b>Peec AI</b>", body_style), Paragraph("Startups, Agencies", body_style), Paragraph("$95 - $245/mo", body_style), Paragraph("Daily prompt runs", body_style), Paragraph("Clean UI, unlimited seats, AI Shopping Analytics", body_style)],
        [Paragraph("<b>ZipTie.dev</b>", body_style), Paragraph("SEO Teams", body_style), Paragraph("~$79/mo", body_style), Paragraph("Real UI crawling", body_style), Paragraph("Simulates user UI behaviors, auto query generation", body_style)],
        [Paragraph("<b>Otterly.AI</b>", body_style), Paragraph("Small Business", body_style), Paragraph("~$29/mo", body_style), Paragraph("Multi-model monitor", body_style), Paragraph("Low cost, quick setup, simple brand alerts", body_style)],
        [Paragraph("<b>Scrunch AI</b>", body_style), Paragraph("Data Analysts", body_style), Paragraph("$83 - $417+/mo", body_style), Paragraph("Cohort engine", body_style), Paragraph("Cohort breakdowns, metric customization", body_style)],
        [Paragraph("<b>Hall</b>", body_style), Paragraph("Brand & PR", body_style), Paragraph("$199 - $499+/mo", body_style), Paragraph("Continuous query", body_style), Paragraph("Sentiment tracking, citation drift, crawler activity", body_style)]
    ]

    pdf_table = Table(table_raw_data, colWidths=[1.1*inch, 1.2*inch, 1.0*inch, 1.2*inch, 2.5*inch])
    pdf_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#1A365D")),
        ('TEXTCOLOR', (0,0), (-1,0), colors.white),
        ('ALIGN', (0,0), (-1,-1), 'LEFT'),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#CBD5E0")),
        ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.HexColor("#F7FAFC"), colors.HexColor("#EDF2F7")]),
        ('BOTTOMPADDING', (0,0), (-1,-1), 4),
        ('TOPPADDING', (0,0), (-1,-1), 4),
    ]))
    story.append(pdf_table)

    # Section 5
    story.append(Paragraph("5. Technical Analysis of VerifiedDR", h1_style))
    story.append(Paragraph("Combines on-page technical diagnostics (local 47-check browser audit) with multi-engine query tracking to identify indexing blockages and generate developer-ready prompts for AI workspaces.", body_style))
    story.append(Paragraph("<b>1. AI Crawler Access:</b> Inspects robots.txt for GPTBot, OAI-SearchBot, ClaudeBot, PerplexityBot.", bullet_style))
    story.append(Paragraph("<b>2. JavaScript Blindness:</b> Compares raw HTML vs. rendered DOM to prevent invisible JS content.", bullet_style))
    story.append(Paragraph("<b>3. Structured Data:</b> Validates JSON-LD schema (Organization, Product, FAQPage, Article).", bullet_style))
    story.append(Paragraph("<b>4. llms.txt Verification:</b> Confirms root markdown file availability.", bullet_style))
    story.append(Paragraph("<b>5. TrueDR Score:</b> Cross-references Domain Rating with verified search traffic and trust signals.", bullet_style))

    # Section 6 Checklist
    story.append(Paragraph("6. Optimization & Technical Checklists", h1_style))
    story.append(Paragraph("<b>[ ] Include Expert Quotes & Stats:</b> Boosts PAWC citation visibility by +31% to +41%.", bullet_style))
    story.append(Paragraph("<b>[ ] Unblock AI Crawlers:</b> Explicitly permit OAI-SearchBot, PerplexityBot, and ClaudeBot in robots.txt.", bullet_style))
    story.append(Paragraph("<b>[ ] Host llms.txt:</b> Create clear markdown overview at domain root domain.com/llms.txt.", bullet_style))
    story.append(Paragraph("<b>[ ] Eliminate Keyword Stuffing:</b> Avoid artificial repetition (-8.0% penalty).", bullet_style))

    # Section 7 Roadmap
    story.append(Paragraph("7. Strategic Implementation Roadmap", h1_style))
    story.append(Paragraph("<b>• Phase 1 (Days 1–14):</b> Audit robots.txt, fix JS blindness, deploy llms.txt, validate JSON-LD.", bullet_style))
    story.append(Paragraph("<b>• Phase 2 (Days 15–45):</b> Add expert quotes/stats, restructure with direct-answer openings.", bullet_style))
    story.append(Paragraph("<b>• Phase 3 (Ongoing):</b> Monitor neutral prompts and generate developer prompts for instant patches.", bullet_style))

    doc.build(story)
    print(f"PDF saved to {pdf_path}")

if __name__ == "__main__":
    create_word_docx()
    create_pdf()
